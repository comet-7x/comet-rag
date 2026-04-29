from typing import Any, TypeVar, Generic

from pydantic import BaseModel
from docx.document import Document as DocumentObject

T = TypeVar("T")


class BaseDocument(BaseModel, Generic[T]):
    model_config = {"arbitrary_types_allowed": True}
    elements: T
    metadata: dict[str, Any]


class ByteDocument(BaseDocument[bytes]):
    pass


class DocxDocument(BaseDocument[DocumentObject]):
    pass
