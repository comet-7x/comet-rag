"""
Production-grade DOCX parser for Comet-RAG.

Block types emitted
-------------------
text      — body paragraph            {type, content, style}
heading   — heading paragraph         {type, content, level, is_numbered[, number]}
list      — ordered / unordered list  {type, attribute, ilevel, content: [Block, …]}
table     — table block               {type, content, rows, row_count, col_count}
image     — embedded image            {type, content, format, name, alt_text, id, width_px, height_px}
equation  — standalone equation       {type, content}   (LaTeX or raw text)
caption   — figure / table caption    {type, content}
header    — page header               {type, content}
footer    — page footer               {type, content}

Inline rich text in ``content`` uses Markdown conventions:
  **bold**  *italic*  ***bold+italic***  ~~strikethrough~~  [text](url)
  Inline equations:   $latex$
  Standalone equations are wrapped in $$…$$.

OMML → LaTeX conversion is handled by the built-in ``_omml`` module
(no external dependencies beyond lxml and loguru).
"""

import asyncio
import base64
import re
from collections.abc import Iterator
from dataclasses import dataclass
from typing import Any

from docx.document import Document as DocumentObject
from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import qn
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from loguru import logger
from lxml import etree  # pyright: ignore[reportAttributeAccessIssue]

from comet_rag.engines.converters.types import DocxDocument
from comet_rag.engines.parsers.base_parser import BaseParser
from comet_rag.engines.parsers.docx_parser.omml import oMath2Latex as _oMath2Latex
from comet_rag.engines.parsers.types import Block, DocxParsedContent

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

_R_EMBED = f"{{{_R}}}embed"
_XML_VAL = f"{{{_W}}}val"
_EMU_PER_PX = 9525  # 914 400 EMU/inch ÷ 96 DPI

#: 单张表允许展开的单元格数上限，超出即跳过并留占位（见 `_skip_oversized_table`）。
#:
#: 这里用常数是**恰当的**：它是资源预算，不是语义上界 —— 与"缺列数不得超过
#: 表格网格宽度"那种从文档本身推出来的约束不是一回事。100 万格约合 8 MB 指针，
#: 而现实里最大的表（几百列 × 几千行）也远够不着。
_MAX_TABLE_CELLS = 1_000_000

# Inline wrapper tags treated as transparent pass-throughs
_TRANSPARENT_INLINE: frozenset[str] = frozenset(
    {"bdo", "customXml", "dir", "fldSimple", "ins", "moveTo", "smartTag"}
)


@dataclass(frozen=True, eq=True)
class _Fmt:
    bold: bool = False
    italic: bool = False
    underline: bool = False
    strikethrough: bool = False


_PLAIN = _Fmt()

# (text, formatting, hyperlink-url | None)
_PElem = tuple[str, _Fmt, str | None]


@dataclass(frozen=True)
class _EqSeg:
    latex: str


_Seg = _PElem | _EqSeg


def _merge_segs(elems: list[_Seg]) -> list[_Seg]:
    """Merge adjacent text segments that share identical format and URL."""
    merged: list[_Seg] = []
    for seg in elems:
        if (
            merged
            and not isinstance(seg, _EqSeg)
            and not isinstance(merged[-1], _EqSeg)
            and seg[1] == merged[-1][1]
            and seg[2] == merged[-1][2]
        ):
            prev = merged[-1]
            merged[-1] = (prev[0] + seg[0], prev[1], prev[2])
        else:
            merged.append(seg)
    return merged


def _qname(element: Any) -> str:
    if not isinstance(getattr(element, "tag", None), str):
        return ""
    return etree.QName(element).localname


def _apply_fmt_and_url(text: str, fmt: _Fmt, url: str | None) -> str:
    """Wrap *text* with Markdown formatting markers and/or a hyperlink."""
    stripped = text.strip()
    if not stripped:
        return text

    prefix = text[: len(text) - len(text.lstrip())]
    suffix = text[len(text.rstrip()) :]

    inner = stripped
    if fmt.bold and fmt.italic:
        inner = f"***{inner}***"
    elif fmt.bold:
        inner = f"**{inner}**"
    elif fmt.italic:
        inner = f"*{inner}*"
    if fmt.strikethrough:
        inner = f"~~{inner}~~"

    if url and url.strip() not in ("", "."):
        url_esc = url.replace("(", "%28").replace(")", "%29")
        return f"{prefix}[{inner}]({url_esc}){suffix}"

    return f"{prefix}{inner}{suffix}"


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    escaped_rows = [
        [cell.replace("|", "\\|").replace("\n", "<br>") for cell in row] for row in rows
    ]
    header, *body = escaped_rows
    sep = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
        *("| " + " | ".join(row) + " |" for row in body),
    ]
    return "\n".join(lines)


def _is_hidden_run(run: Run) -> bool:
    rpr = run._element.find(f"{{{_W}}}rPr")
    if rpr is None:
        return False
    return (
        rpr.find(f"{{{_W}}}webHidden") is not None
        or rpr.find(f"{{{_W}}}vanish") is not None
    )


def _resolve_style_bool(style_obj: Any, attr: str) -> bool | None:
    """Walk the style inheritance chain and resolve a boolean font attribute."""
    style = style_obj
    while style is not None:
        font = getattr(style, "font", None)
        if font is not None:
            if attr == "underline":
                value = font.underline
            elif attr == "strikethrough":
                value = font.strike
            else:
                value = getattr(font, attr, None)
            if value is not None:
                return bool(value)
        style = getattr(style, "base_style", None)
    return None


def _get_run_fmt(run: Run) -> _Fmt:
    """Resolve run formatting with full style-inheritance chain."""

    def resolve(attr: str) -> bool:
        direct = (
            run.underline
            if attr == "underline"
            else (
                run.font.strike if attr == "strikethrough" else getattr(run, attr, None)
            )
        )
        if direct is not None:
            return bool(direct)

        # Character-style chain — skip Hyperlink default style to avoid
        # spurious underlines on every hyperlink run.
        run_style = getattr(run, "style", None)
        sid = str(getattr(run_style, "style_id", "") or "").lower()
        sname = str(getattr(run_style, "name", "") or "").lower()
        if "hyperlink" not in sid and "hyperlink" not in sname:
            v = _resolve_style_bool(run_style, attr)
            if v is not None:
                return v

        # Paragraph-style chain
        parent = getattr(run, "_parent", None)
        v = _resolve_style_bool(getattr(parent, "style", None), attr)
        if v is not None:
            return v
        return False

    underline = resolve("underline")
    # CJK emphasis mark (w:em) also renders as an underline-like decoration
    rpr = run._element.find(f"{{{_W}}}rPr")
    if rpr is not None:
        em = rpr.find(f"{{{_W}}}em")
        if em is not None and em.get(_XML_VAL, "") not in ("", "none"):
            underline = True

    return _Fmt(
        bold=resolve("bold"),
        italic=resolve("italic"),
        underline=underline,
        strikethrough=resolve("strikethrough"),
    )


class DocxParser(BaseParser[DocxDocument, DocxParsedContent]):
    """Parse a DocxDocument into a list of semantically typed blocks."""

    def __init__(
        self,
        *,
        heading_numbers: bool = False,
        max_table_cells: int = _MAX_TABLE_CELLS,
    ) -> None:
        self._max_table_cells = max_table_cells
        self._doc: DocumentObject | None = None
        self._doc_part: Any = None
        self._blocks: list[Block] = []
        # Active list tracking
        self._pre_num_id: int = -1
        self._pre_ilevel: int = -1
        self._list_stack: list[Block] = []
        # Heading auto-number extraction
        self._heading_numbers = heading_numbers
        self._h_counters: list[int] = [0] * 9
        self._h_numId: int | None = None
        self._h_fmts: list[str] = []
        self._h_starts: list[int] = []
        self._styles_root: Any = None
        self._footnotes: dict[int, str] | None = None

    def parse(self, document: DocxDocument) -> DocxParsedContent:
        self._doc = document.elements
        self._doc_part = self._doc.part
        self._blocks = []
        self._pre_num_id = -1
        self._pre_ilevel = -1
        self._list_stack = []
        self._h_counters = [0] * 9
        self._h_numId = None
        self._h_fmts = []
        self._h_starts = []
        self._styles_root = None
        self._footnotes = None

        # python-docx 的 `Document.element` 没有返回标注，静态上退化成
        # `BaseOxmlElement`，看不见 `body`。运行时它是 `CT_Document`。
        self._walk(self._doc.element.body)  # pyright: ignore[reportAttributeAccessIssue]
        self._add_headers_footers()

        return DocxParsedContent(blocks=self._blocks, metadata=document.metadata)

    async def aparse(self, document: DocxDocument) -> DocxParsedContent:
        """在线程中执行 python-docx/lxml 的同步解析，避免阻塞事件循环。"""
        return await asyncio.to_thread(self.parse, document)

    def _walk(self, container: Any) -> None:
        for _, child in enumerate(container):
            tag = _qname(child)
            if tag == "p":
                self._handle_paragraph(child)
            elif tag == "tbl":
                self._close_list()
                self._handle_table(child)
            elif tag == "sdt":
                sdt_content = child.find(f"{{{_W}}}sdtContent")
                if sdt_content is not None:
                    self._walk(sdt_content)
            elif tag in _TRANSPARENT_INLINE:
                self._walk(child)

    def _handle_paragraph(self, element: Any) -> None:  # noqa: C901 (complexity ok here)
        paragraph = Paragraph(element, self._doc)  # pyright: ignore[reportArgumentType]
        p_elems = self._get_paragraph_elements(paragraph)
        eq_segs = [s for s in p_elems if isinstance(s, _EqSeg)]
        has_plain_text = any(s[0].strip() for s in p_elems if not isinstance(s, _EqSeg))
        rich_text = self._build_rich_text(p_elems)

        style_name = (paragraph.style.name if paragraph.style else None) or "Normal"
        h_level = self._get_heading_level(paragraph)
        numid, ilevel = self._get_numId_ilvl(paragraph)
        if numid == 0:
            numid = None

        # Extract images that live inside this paragraph element
        image_blocks = self._extract_images(element)

        # List item (non-heading)
        if numid is not None and ilevel is not None and h_level is None:
            if rich_text:
                self._add_list_item(
                    numid, ilevel, rich_text, self._is_numbered_list(numid, ilevel)
                )
            self._blocks.extend(image_blocks)
            return

        # All non-list branches close any active list first
        self._close_list()

        # Heading
        if h_level is not None:
            if rich_text:
                h_numid = self._resolve_heading_numid(element)
                block: Block = {
                    "type": "heading",
                    "level": h_level,
                    "content": rich_text,
                    "is_numbered": h_numid is not None,
                }
                if self._heading_numbers and h_numid is not None:
                    num_str = self._compute_heading_number(h_numid, h_level)
                    if num_str:
                        block["number"] = num_str
                self._blocks.append(block)

        # Standalone equation
        elif eq_segs and not has_plain_text:
            eq_content = " ".join(s.latex for s in eq_segs).strip()
            if eq_content:
                self._blocks.append(
                    {"type": "equation", "content": f"$${eq_content}$$"}
                )

        # Caption (SEQ-field based detection)
        elif self._is_caption(element):
            if rich_text:
                self._blocks.append({"type": "caption", "content": rich_text})

        # Body text
        elif rich_text:
            self._blocks.append(
                {"type": "text", "content": rich_text, "style": style_name}
            )

        self._blocks.extend(image_blocks)

    # Inner-content iterator
    # Handles inline w:sdt content controls and transparent wrappers that
    # python-docx's iter_inner_content() does not traverse.
    def _iter_inner_content(
        self,
        paragraph: Paragraph,
        container: Any | None = None,
    ) -> Iterator[Run | Hyperlink | Any]:
        if container is None:
            container = paragraph._element
        for child in container:
            tag = _qname(child)
            if tag == "r":
                yield Run(child, paragraph)
            elif tag == "hyperlink":
                yield Hyperlink(child, paragraph)
            elif tag == "oMath":
                yield child
            elif tag == "oMathPara":
                yield from self._iter_inner_content(paragraph, child)
            elif tag == "sdt":
                sdt_content = child.find(f"{{{_W}}}sdtContent")
                if sdt_content is not None:
                    yield from self._iter_inner_content(paragraph, sdt_content)
            elif tag in _TRANSPARENT_INLINE:
                yield from self._iter_inner_content(paragraph, child)

    def _get_paragraph_elements(self, paragraph: Paragraph) -> list[_Seg]:  # noqa: C901
        """
        Return an ordered list of text and equation segments for the paragraph.

        Handles:
        - w:hyperlink wrappers (external URLs)
        - w:fldChar / w:instrText field-code hyperlinks
        - Hidden runs (w:vanish, w:webHidden)
        - m:oMath / m:oMathPara inline and block equations
        - Consecutive same-format runs are merged for compactness
        """
        elems: list[_Seg] = []

        # Field-code hyperlink tracking state
        _field_in = False
        _field_phase: str | None = None  # "instr" | "result"
        _field_url: str | None = None
        _field_text = ""
        _field_fmt: _Fmt | None = None

        # Current accumulation group (merges consecutive same-format runs)
        _grp_text = ""
        _grp_fmt: _Fmt | None = None

        def _flush() -> None:
            nonlocal _grp_text, _grp_fmt
            if _grp_text:
                elems.append((_grp_text, _grp_fmt or _PLAIN, None))
            _grp_text = ""
            _grp_fmt = None

        for c in self._iter_inner_content(paragraph):
            # Equation element
            if not isinstance(c, (Run, Hyperlink)):
                _flush()
                latex = self._convert_omath(c)
                if latex:
                    elems.append(_EqSeg(latex=latex))
                continue

            # Hyperlink element
            if isinstance(c, Hyperlink):
                address = c.address
                # Keep only external URLs; skip internal anchors (TOC noise)
                url = address if (address and "://" in address) else None
                _flush()
                for h_run in c.runs:
                    if _is_hidden_run(h_run):
                        continue
                    h_text = h_run.text or ""
                    if h_text:
                        elems.append((h_text, _get_run_fmt(h_run), url))
                continue

            if not isinstance(c, Run):
                continue

            # fldChar (field boundary marker)
            fld_char = c._element.find(f"{{{_W}}}fldChar")
            if fld_char is not None:
                fld_type = fld_char.get(f"{{{_W}}}fldCharType")
                if fld_type == "begin":
                    _flush()
                    _field_in = True
                    _field_phase = "instr"
                    _field_url = _field_text = ""
                    _field_fmt = None
                elif fld_type == "separate":
                    _field_phase = "result"
                elif fld_type == "end":
                    if _field_text.strip():
                        elems.append(
                            (_field_text, _field_fmt or _PLAIN, _field_url or None)
                        )
                    _field_in = False
                    _field_phase = _field_url = None
                    _field_text = ""
                    _field_fmt = None
                continue  # fldChar runs carry no displayable text

            # instrText: extract HYPERLINK url
            instr = c._element.find(f"{{{_W}}}instrText")
            if instr is not None:
                if _field_phase == "instr" and instr.text:
                    m = re.search(r'HYPERLINK\s+"([^"]+)"', instr.text)
                    if m:
                        _field_url = m.group(1)
                continue  # never displayable

            # Skip non-display runs inside an instr block
            if _field_in and _field_phase == "instr":
                continue

            # Accumulate field result display text
            if _field_in and _field_phase == "result":
                if c._element.find(f"{{{_W}}}t") is not None:
                    _field_text += c.text or ""
                    if _field_fmt is None:
                        _field_fmt = _get_run_fmt(c)
                continue

            # Hidden run
            if _is_hidden_run(c):
                continue

            # Footnote reference
            fn_ref = c._element.find(f"{{{_W}}}footnoteReference")
            if fn_ref is not None:
                fn_id = int(fn_ref.get(f"{{{_W}}}id", -1))
                fn_text = self._load_footnotes().get(fn_id, "")
                if fn_text:
                    _flush()
                    elems.append((f"（{fn_text}）", _PLAIN, None))
                continue

            # Normal run
            text = c.text or ""
            if not text:
                continue
            fmt = _get_run_fmt(c)

            if fmt == _grp_fmt or _grp_fmt is None:
                _grp_fmt = fmt
                _grp_text += text
            else:
                _flush()
                _grp_fmt = fmt
                _grp_text = text

        _flush()
        return _merge_segs(elems)

    def _build_rich_text(self, elems: list[_Seg]) -> str:
        return "".join(
            f"${seg.latex}$"
            if isinstance(seg, _EqSeg)
            else _apply_fmt_and_url(seg[0], seg[1], seg[2])
            for seg in elems
        )

    def _convert_omath(self, element: Any) -> str:
        """Convert an oMath element to LaTeX, falling back to raw text on error."""
        try:
            return str(_oMath2Latex(element)).strip()
        except Exception as exc:
            logger.debug(f"oMath2Latex failed: {exc}")
            return "".join(
                node.text
                for node in element.iter()
                if _qname(node) == "t" and node.text
            )

    def _extract_images(self, para_element: Any) -> list[Block]:
        blocks: list[Block] = []
        for drawing in para_element.iter(qn("w:drawing")):
            docPr = drawing.find(f".//{qn('wp:docPr')}")
            name = docPr.get("name", "") if docPr is not None else ""
            alt_text = docPr.get("descr", "") if docPr is not None else ""
            docPr_id = docPr.get("id", "") if docPr is not None else ""

            extent = drawing.find(f".//{qn('wp:extent')}")
            width_px = 0
            height_px = 0
            if extent is not None:
                try:
                    width_px = int(extent.get("cx", 0)) // _EMU_PER_PX
                    height_px = int(extent.get("cy", 0)) // _EMU_PER_PX
                except (ValueError, TypeError):
                    pass

            blip = drawing.find(f".//{qn('a:blip')}")
            if blip is None:
                continue
            rId = blip.get(_R_EMBED)
            if rId not in self._doc_part.rels:
                continue

            image_part = self._doc_part.rels[rId].target_part
            blocks.append(
                {
                    "type": "image",
                    "content": base64.b64encode(image_part.blob).decode(),
                    "format": image_part.content_type.split("/")[-1],
                    "name": name,
                    "alt_text": alt_text,
                    "id": docPr_id,
                    "width_px": width_px,
                    "height_px": height_px,
                }
            )
        return blocks

    def _cell_to_text(self, cell: Any) -> str:
        """Extract rich text from a table cell, including inline equations."""
        parts: list[str] = []
        for para in cell.paragraphs:
            elems = self._get_paragraph_elements(para)
            text = self._build_rich_text(elems).strip()
            if text:
                parts.append(text)
        return "\n".join(parts)

    @staticmethod
    def _raw_gaps(row: Any) -> tuple[int, int]:
        """一行两端声明的缺列数，未封顶。

        这两个值直接来自文档 XML —— 在"用户上传文件"这条路上等于**攻击者
        可控的整数**，所以取出来之后必须封顶（见 `_handle_table`），
        不能直接拿去 `[""] * n`。
        """
        return (
            max(int(getattr(row, "grid_cols_before", 0) or 0), 0),
            max(int(getattr(row, "grid_cols_after", 0) or 0), 0),
        )

    @staticmethod
    def _declared_spans(row: Any) -> list[tuple[int, bool]]:
        """本行每个 `tc` 的 `(自己声明的 span, 是不是纵向合并的续格)`。

        续格的判定：`w:vMerge` 存在且 `w:val` 不是 `restart`（OOXML 里省略
        `val` 就等于 `continue`）。这类格子**自己的 `gridSpan` 不作数** ——
        展开时用的是上方根单元格的，见 `_projected_cells()`。
        """
        out: list[tuple[int, bool]] = []
        for tc in row._tr.tc_lst:  # noqa: SLF001
            tc_pr = tc.tcPr
            merge = None if tc_pr is None else tc_pr.vMerge
            out.append((max(tc.grid_span, 1), merge is not None and merge != "restart"))
        return out

    def _row_extent(self, row: Any, grid_width: int) -> tuple[int, int]:
        """`(两端缺列数之和, 本行 tc 自己声明的 span 之和)` —— **不碰 `row.cells`**。

        为什么不能碰：`row.cells` 是按网格列展开的，长度由文档里的 `gridSpan`
        决定。一份 2 列的表里塞一个 `gridSpan=10000000`，python-docx 就会建出
        一千万个 `_Cell`（实测 75 MB）—— 而这发生在解析循环**之前**，循环里
        再怎么设防都拦不住（PR #34 评审）。

        所以只能从 XML 层数：`tc` 的个数与各自的 span 都是逐个元素声明的，
        规模等于文件规模，攻击者拿不到杠杆。

        这里走 `row._tr` / `tc.grid_span` 这条内部路径是**刻意的**，与
        `_row_to_cells()` 里坚持用公开 `cell.grid_span` 不矛盾：那边是在读
        已经安全展开的数据，这边是在决定要不要展开 —— 而公开 API 恰恰正是
        那个无界的东西，用它做准入检查等于先中招再判断。
        """
        before, after = self._raw_gaps(row)
        spans = sum(max(tc.grid_span, 1) for tc in row._tr.tc_lst)  # noqa: SLF001
        return min(before, grid_width) + min(after, grid_width), spans

    def _projected_cells(self, table: Any, grid_width: int) -> int:
        """整表展开后的单元格数**上界**，全程不碰 `row.cells`。

        ## 为什么不能只看本行自己声明的 span

        `vMerge="continue"` 的续格会继承**上方根单元格**的 `gridSpan`，而它
        自己的 `tc` 往往省略 `gridSpan`（本地值就是 1）。实测：

            r0: XML 里 grid_span=[5]   row.cells 实际长度 = 5   (vMerge restart)
            r1: XML 里 grid_span=[1]   row.cells 实际长度 = 5   (vMerge continue)

        于是"逐行累加本地 span"会严重低估：根行声明 `gridSpan=S`、后面跟 R 个
        续行，本地和约 `S + R`，实际展开却是 `S × (R+1)`。取 S=999000、R=999
        就能把投影压在默认上限之下，而真实展开接近十亿格（PR #34 评审）。

        ## 上界只能逐格取，不能按"历史最大行宽"取

        我先前用的是"第 i 行宽度 ≤ 两端缺列 + max(前 i 行各自的 span 之和)"。
        **那不是上界**：同一行可以既继承一个宽 span、又自己再声明一个宽 span。
        实测（PR #34 评审给出的反例）：

            r0: 本地 span=[5]     row.cells 实际 = 5
            r1: 本地 span=[1, 5]  row.cells 实际 = 10   ← 继承 5 + 新声明 5

            按行取 max 的投影 = 11 < 实际 15

        改成**逐格**取：

            续格   的有效 span ≤ 见过的最大单格 span（它继承的那个根，
                                 一定是某处声明过的一个 gridSpan）
            非续格 的有效 span  = 它自己声明的 gridSpan

        这一条不依赖 python-docx 如何解析重叠布局，纯粹是"每格的展开宽度都
        来自某个声明过的 gridSpan"，所以是硬上界。

        非续格用真实值而不是一律取 max，是为了不误伤合法文档：一张 50 列的表
        若某行是整行合并的表头，"全部按 50 算"会把正文行也放大 50 倍。实测
        4 份真实文档 17 张表，最大投影 144（预算 100 万），没有误拒风险。

        ## `widest_single` 刻意先更新再计分

        于是本行**后面**才声明的宽格，也会被用来估本行前面的续格。严格说
        续格只可能继承**更上方**的根，所以这是过估计 —— 但刻意保留（PR #34
        评审也确认过安全性没问题）：它顺带兜住了"续格上方根本没有根"这类
        畸形输入，而先计分再更新反倒会在首行给出 0。宁可松，不可漏。
        """
        projected = 0
        widest_single = 0
        for row in table.rows:
            gaps, _ = self._row_extent(row, grid_width)
            spans = self._declared_spans(row)
            widest_single = max([widest_single, *(s for s, _ in spans)])
            projected += gaps + sum(
                widest_single if inherited else span for span, inherited in spans
            )
            if projected > self._max_table_cells:
                break  # 已经超了，没必要把剩下的行也数一遍
        return projected

    def _row_to_cells(self, row: Any, before: int, after: int) -> list[str]:
        """一行的单元格文本，横向合并（gridSpan）的续格留空。

        ## 判据必须来自结构，不能是文本相等

        python-docx 的 `row.cells` 按**网格列**返回：一个横跨两列的单元格会被
        吐出两次。`grid_span` 直接告诉我们它横跨几列 —— "这一格后面还有几个
        位置是它的续格"这件事，结构里写得明明白白。

        此前用的是"相邻文本相等就折叠"。那等于拿数据去猜结构，而相邻两列取值
        相同在真实表格里再常见不过：

            文档里是 3 列  [季度, Q1, Q1] / [营收, 100, 100]
            解析出的是     [季度, Q1]     / [营收, 100]

        整整一列凭空消失，**不报错也不打日志**，一路进到向量库（#18）。
        两个相邻的空单元格同样会被折叠成一个。

        ## 用 `grid_span` 而不是比较底层的 `tc` 对象

        两者都能识别出续格（同一个单元格会被吐出多次，`_tc` 自然相同），
        但 `_tc` 是 python-docx 的**私有属性**，而 `grid_span` 是公开且有
        文档的 API。依赖内部实现的代价不是抽象的：`python-docx>=1.2.0` 没有
        上限，哪次升级把它改掉，这里要么当场 AttributeError，要么更糟 ——
        识别不出合并却继续静默出错（PR #32 评审）。

        ## 续格留空，而不是删掉

        Markdown 表格没有 colspan，续格只能空着。但删掉会让它后面的列整体左移：

            正确  [合并单元格, "",  C]
            错误  [合并单元格, C,  ""]      ← C 落到了第 1 列

        补齐是在**行尾**做的，所以出现在行中间的合并会把后面所有列都错位。
        留空则天然对齐，`col_count` 也才是真实的网格宽度。

        ## 两端的"缺列"也要占位（gridBefore / gridAfter）

        Word 允许一行**晚开始**或**早结束**：`w:gridBefore` / `w:gridAfter`
        声明该行头尾各有几个网格列压根不存在（常见于缩进的子表格行、
        跨页表格的续行）。python-docx 把它们暴露为 `grid_cols_before` /
        `grid_cols_after`，且明确写着"these are not simply empty cells"。

        不补的话，"晚开始"那行的所有单元格都会左移一格：

            正确  [A, B, C] / ["", y, z]
            错误  [A, B, C] / [y, z, ""]      ← y 落到了第 0 列

        与漏掉合并续格是**同一类错位**，只是成因不同（#33）。

        ## 纵向合并（vMerge）不在此列

        它的 `grid_span` 是 1，文本会沿列重复出现 —— 本方法按行处理，碰不到
        它，行为与此前一致。对检索而言每行自带上下文反而是好事。
        """
        # 行首的缺列：不是空单元格，是不存在的网格列。`before` / `after`
        # 已由调用方按网格宽度封顶，这里不再重复校验。
        cells: list[str] = [""] * before
        continuations = 0
        for cell in row.cells:
            if continuations:
                cells.append("")
                continuations -= 1
                continue
            cells.append(self._cell_to_text(cell))
            # 横跨 n 列 ⇒ 后面 n-1 个网格位置是同一个单元格的续格
            continuations = max(cell.grid_span - 1, 0)
        cells.extend([""] * after)
        return cells

    def _skip_oversized_table(self, rows: int, cells: int) -> None:
        """超预算的表：不展开，但**在原位留一条能被检索到的说明**。

        为什么不截断成前 N 格：那会产出一张**看起来完整、实则残缺**的表。
        检索到它的人无从判断后面还有内容，于是把半份数据当成全部 —— 这正是
        #18（整列消失）与 #33（整行错位）的同一个失败模式，本不该在处理
        它们的过程中再造一个。

        为什么不干脆丢掉：那样知识库里完全看不出这里曾经有过东西，只有运维
        翻日志才知道。留占位则让"缺失"本身可见、可检索。

        用 `caption` 类型是因为它已经在 `DocxParsedContent.text` 的白名单里 ——
        占位块必须能进入正文，否则切块与向量化都会把它丢掉，也就搜不到了。
        """
        logger.warning(
            f"表格 {rows} 行、展开后约 {cells:,} 个单元格，超出上限 "
            f"{self._max_table_cells:,}，已跳过并留占位 —— "
            f"可用 DocxParser(max_table_cells=…) 调整"
        )
        self._blocks.append(
            {
                "type": "caption",
                "content": f"[表格未收录：{rows} 行，约 {cells:,} 个单元格，"
                f"超出解析上限 {self._max_table_cells:,}]",
            }
        )

    def _handle_table(self, element: Any) -> None:
        from docx.table import Table

        table = Table(element, self._doc)  # pyright: ignore[reportArgumentType]
        try:
            # 声明的网格宽度，用作单行缺列数的上界
            grid_width = len(table.columns)
        except InvalidXmlError:
            # `tblGrid` 是 schema 要求的必需元素。缺了它就无从确定网格宽度 ——
            # 这不是"某一行有问题"，是整张表不合法。
            logger.warning("表格缺少 <w:tblGrid>，无从确定网格宽度，已跳过该表")
            return

        # ── 先定界，再碰任何一行 ───────────────────────────────────────────
        #
        # 展开后的规模由**三个**文档可控的量决定，任何一个不设防都能放大：
        #
        #   gridBefore / gridAfter   一行头尾声明缺几列
        #   gridSpan                 一格横跨几列
        #   行数 × 网格宽度           整表的矩形大小
        #
        # 实测（均为 PR #34 评审）：
        #   · N=R=3000 的"宽网格 + 短行"：约 9,000 个 XML 元素、38 KB 的 docx，
        #     展开成 900 万个单元格、峰值 87 MB
        #   · 2 列的表里塞一个 gridSpan=10000000：`row.cells` 直接返回一千万个
        #     `_Cell`，光是建它就 75 MB
        #
        # 所以判据必须**逐个 `tc` 从 XML 上数**，而且要在碰 `row.cells` 之前
        # 算完 —— 后者的长度正是由 gridSpan 决定的，等它建好就晚了。
        #
        # 还有第四个维度：`vMerge` 的续格会继承上方根单元格的 span，见
        # `_projected_cells()`。
        projected = self._projected_cells(table, grid_width)
        if projected > self._max_table_cells:
            self._skip_oversized_table(len(table.rows), projected)
            return

        # ── 到这里，规模已经有界，可以逐行展开 ─────────────────────────────
        cleaned: list[list[str]] = []
        truncated_rows = 0
        largest_gap = 0
        for row in table.rows:
            raw_before, raw_after = self._raw_gaps(row)
            before, after = min(raw_before, grid_width), min(raw_after, grid_width)
            if (before, after) != (raw_before, raw_after):
                truncated_rows += 1
                largest_gap = max(largest_gap, raw_before, raw_after)
            cleaned.append(self._row_to_cells(row, before, after))

        if truncated_rows:
            # 每张表只报一条。按行报的话，一份畸形文档能直接产出成千上万行
            # 日志 —— 把"文档内容"放大成"日志写入量"，与上面那个内存放大
            # 是同一类问题（PR #34 评审）。
            logger.warning(
                f"表格有 {truncated_rows} 行声明的缺列数超出网格宽度 {grid_width}"
                f"（最大 {largest_gap}），已截断 —— 文档很可能是畸形的"
            )

        if not any(any(r) for r in cleaned):
            return

        # 两端的缺列补上之后（见 `_row_to_cells`），每行的宽度**本就应当**等于
        # 表格的网格宽度 —— 实测 17 张真实表格无一例外。走到这里还需要补，
        # 说明文档里的网格声明本身不自洽。
        #
        # 补是为了不让下游拿到参差的行，但必须留痕：这一步只会往**行尾**填，
        # 而真正缺的列可能在中间；静默做等于把错位藏起来（#33）。
        widths = {len(r) for r in cleaned}
        max_cols = max(widths)
        if len(widths) > 1:
            # 只报区间与种数，不列全集：这条日志由**文档内容**触发，
            # 而不同宽度的数量上限是行数 —— 一份畸形文档能让它长到几万项，
            # 把日志管道撑坏（PR #34 评审）。
            logger.warning(
                f"表格各行的网格宽度不一致（{len(widths)} 种，"
                f"{min(widths)}–{max_cols} 列），已按最宽的 {max_cols} 列"
                f"在行尾补齐 —— 该表的列对应关系可能不准"
            )
        for row in cleaned:
            row.extend("" for _ in range(max_cols - len(row)))

        self._blocks.append(
            {
                "type": "table",
                "content": _table_to_markdown(cleaned),
                "rows": cleaned,
                "row_count": len(cleaned),
                "col_count": max_cols,
            }
        )

    def _get_numId_ilvl(self, paragraph: Paragraph) -> tuple[int | None, int | None]:
        # 必须剔掉 `None` 键（默认命名空间）：`find()` 走的是 ElementPath，
        # 它不支持空前缀，喂进去会直接抛 ValueError，整份文件的列表编号识别
        # 当场崩掉。
        #
        # 手上的 docx 大多把命名空间都写成带前缀的，所以这个缺陷此前一直没有
        # 暴露 —— 但那是**样本的性质，不是格式的保证**：默认命名空间在
        # OOXML 里完全合法，只是不常见。拿"实际文件通常都带前缀"当不变式，
        # 等于把正确性押在数据分布上。
        nsmap = {
            prefix: uri
            for prefix, uri in paragraph._element.nsmap.items()
            if prefix is not None
        }
        numPr = paragraph._element.find(".//w:numPr", namespaces=nsmap)
        if numPr is None:
            return None, None

        def _int(el: Any) -> int | None:
            if el is None:
                return None
            try:
                return int(el.get(_XML_VAL))
            except (TypeError, ValueError):
                return None

        return (
            _int(numPr.find("w:numId", namespaces=nsmap)),
            _int(numPr.find("w:ilvl", namespaces=nsmap)),
        )

    def _is_numbered_list(self, numId: int, ilvl: int) -> bool:
        """Return True when the list level uses a numeric numbering format."""
        _NUMBERED = {
            "decimal",
            "lowerRoman",
            "upperRoman",
            "lowerLetter",
            "upperLetter",
            "decimalZero",
        }
        try:
            numbering_part = next(
                (p for p in self._doc.part.package.parts if "numbering" in p.partname),  # pyright: ignore[reportOptionalMemberAccess]
                None,
            )
            if numbering_part is None:
                return False

            ns = {"w": _W}
            root = numbering_part.element  # pyright: ignore[reportAttributeAccessIssue]

            num_el = root.find(f".//w:num[@w:numId='{numId}']", ns)
            if num_el is None:
                return False

            abs_id_el = num_el.find(".//w:abstractNumId", ns)
            if abs_id_el is None:
                return False
            abs_id = abs_id_el.get(_XML_VAL)

            abs_num = root.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns)
            if abs_num is None:
                return False

            lvl = abs_num.find(f".//w:lvl[@w:ilvl='{ilvl}']", ns)
            if lvl is None:
                return False

            fmt_el = lvl.find(".//w:numFmt", ns)
            if fmt_el is None:
                return False

            return fmt_el.get(_XML_VAL) in _NUMBERED
        except Exception as exc:
            logger.debug(f"_is_numbered_list: {exc}")
            return False

    def _add_list_item(
        self, numid: int, ilevel: int, content: str, is_ordered: bool
    ) -> None:
        attr = "ordered" if is_ordered else "unordered"
        item: Block = {"type": "text", "content": content}

        # New list sequence (different numid or no active list)
        if self._pre_num_id == -1 or self._pre_num_id != numid:
            if self._pre_num_id != -1:
                self._close_list()
            list_block: Block = {
                "type": "list",
                "attribute": attr,
                "ilevel": ilevel,
                "content": [item],
            }
            self._blocks.append(list_block)
            self._list_stack = [list_block]
            self._pre_num_id = numid
            self._pre_ilevel = ilevel
            return

        # Deeper nesting → open a child list block
        if ilevel > self._pre_ilevel:
            child: Block = {
                "type": "list",
                "attribute": attr,
                "ilevel": ilevel,
                "content": [item],
            }
            self._list_stack[-1]["content"].append(child)
            self._list_stack.append(child)
            self._pre_ilevel = ilevel
            return

        # Shallower nesting → pop back to the matching level
        if ilevel < self._pre_ilevel:
            while len(self._list_stack) > 1 and self._list_stack[-1]["ilevel"] > ilevel:
                self._list_stack.pop()
            self._list_stack[-1]["content"].append(item)
            self._pre_ilevel = ilevel
            return

        # Same level
        self._list_stack[-1]["content"].append(item)

    def _close_list(self) -> None:
        self._pre_num_id = -1
        self._pre_ilevel = -1
        self._list_stack = []

    def _get_styles_root(self) -> Any | None:
        if self._styles_root is not None:
            return self._styles_root
        try:
            part = next(
                (
                    p
                    for p in self._doc.part.package.parts  # pyright: ignore[reportOptionalMemberAccess]
                    if p.partname.endswith("styles.xml")
                ),  # pyright: ignore[reportOptionalMemberAccess]
                None,
            )
            if part is not None:
                self._styles_root = part.element  # pyright: ignore[reportAttributeAccessIssue]
        except Exception as exc:
            logger.debug(f"_get_styles_root: {exc}")
        return self._styles_root

    def _load_footnotes(self) -> dict[int, str]:
        """Lazily parse word/footnotes.xml → {footnote_id: plain_text}."""
        if self._footnotes is not None:
            return self._footnotes
        self._footnotes = {}
        try:
            part = next(
                (
                    p
                    for p in self._doc.part.package.parts  # pyright: ignore[reportOptionalMemberAccess]
                    if p.partname.endswith("footnotes.xml")
                ),  # pyright: ignore[reportOptionalMemberAccess]
                None,
            )
            if part is None:
                return self._footnotes
            # FootnotesPart is a generic Part (no .element); parse blob directly.
            root = etree.fromstring(part.blob)  # pyright: ignore[reportAttributeAccessIssue]
            ns = {"w": _W}
            for fn_el in root.findall("w:footnote", ns):
                fn_id_str = fn_el.get(f"{{{_W}}}id")
                if fn_id_str is None:
                    continue
                fn_id = int(fn_id_str)
                if fn_id <= 0:  # skip separator (-1) and continuation (0) markers
                    continue
                text = "".join(
                    t.text for t in fn_el.iter(f"{{{_W}}}t") if t.text
                ).strip()
                if text:
                    self._footnotes[fn_id] = text
        except Exception as exc:
            logger.debug(f"_load_footnotes: {exc}")
        return self._footnotes

    def _resolve_heading_numid(self, element: Any) -> int | None:
        """
        Return the numId for a heading paragraph, or None if not auto-numbered.

        Checks the paragraph's own <w:pPr>/<w:numPr> first, then walks the
        style inheritance chain (needed when numPr lives in styles.xml rather
        than in the paragraph itself, which is the common case in Word).
        """
        ns = {"w": _W}

        # 1. Paragraph-level override
        numpr = element.find(f"{{{_W}}}pPr/{{{_W}}}numPr")
        if numpr is not None:
            numid_el = numpr.find(f"{{{_W}}}numId")
            if numid_el is not None:
                try:
                    val = int(numid_el.get(_XML_VAL, 0))
                except ValueError:
                    val = 0
                return None if val == 0 else val  # val=0 means "suppress"

        # 2. Style inheritance chain
        pstyle = element.find(f"{{{_W}}}pPr/{{{_W}}}pStyle")
        if pstyle is None:
            return None
        style_id = pstyle.get(_XML_VAL)
        if not style_id:
            return None

        root = self._get_styles_root()
        if root is None:
            return None

        visited: set[str] = set()
        current = style_id
        while current and current not in visited:
            visited.add(current)
            style_el = root.find(f".//w:style[@w:styleId='{current}']", ns)
            if style_el is None:
                break
            numpr = style_el.find(".//w:numPr", ns)
            if numpr is not None:
                numid_el = numpr.find("w:numId", ns)
                if numid_el is not None:
                    try:
                        val = int(numid_el.get(_XML_VAL, 0))
                    except ValueError:
                        val = 0
                    return None if val == 0 else val
            based_on = style_el.find("w:basedOn", ns)
            current = based_on.get(_XML_VAL) if based_on is not None else None

        return None

    def _load_heading_formats(self, num_id: int) -> tuple[list[str], list[int]]:
        """Load (format_strings, start_values) for each level of numId from numbering.xml."""
        try:
            numbering_part = next(
                (p for p in self._doc.part.package.parts if "numbering" in p.partname),  # pyright: ignore[reportOptionalMemberAccess]
                None,
            )
            if numbering_part is None:
                return [], []
            ns = {"w": _W}
            root = numbering_part.element  # pyright: ignore[reportAttributeAccessIssue]

            num_el = root.find(f".//w:num[@w:numId='{num_id}']", ns)
            if num_el is None:
                return [], []
            abs_id_el = num_el.find(".//w:abstractNumId", ns)
            if abs_id_el is None:
                return [], []
            abs_id = abs_id_el.get(_XML_VAL)

            abs_num = root.find(f".//w:abstractNum[@w:abstractNumId='{abs_id}']", ns)
            if abs_num is None:
                return [], []

            fmts: list[str] = []
            starts: list[int] = []
            for lvl_el in sorted(
                abs_num.findall("w:lvl", ns),
                key=lambda e: int(e.get(f"{{{_W}}}ilvl", 0)),
            ):
                lvl_text = lvl_el.find("w:lvlText", ns)
                fmts.append(lvl_text.get(_XML_VAL, "") if lvl_text is not None else "")
                start_el = lvl_el.find("w:start", ns)
                starts.append(
                    int(start_el.get(_XML_VAL, 1)) if start_el is not None else 1
                )
            return fmts, starts
        except Exception as exc:
            logger.debug(f"_load_heading_formats: {exc}")
            return [], []

    def _compute_heading_number(self, num_id: int, h_level: int) -> str:
        """Return the auto-number string (e.g. '1.2.3') for a numbered heading."""
        if self._h_numId != num_id:
            self._h_fmts, self._h_starts = self._load_heading_formats(num_id)
            self._h_numId = num_id
            self._h_counters = [0] * 9
            for i, s in enumerate(self._h_starts[:9]):
                self._h_counters[i] = s - 1

        idx = h_level - 1
        self._h_counters[idx] += 1
        for i in range(idx + 1, 9):
            reset_to = self._h_starts[i] - 1 if i < len(self._h_starts) else 0
            self._h_counters[i] = reset_to

        if idx < len(self._h_fmts):
            fmt = self._h_fmts[idx]
            for i in range(9, 0, -1):
                fmt = fmt.replace(f"%{i}", str(self._h_counters[i - 1]))
            return fmt

        return ".".join(str(self._h_counters[i]) for i in range(h_level))

    def _get_heading_level(self, paragraph: Paragraph) -> int | None:
        """
        Walk the paragraph's style inheritance chain.
        Return a 1-based heading level, or None if the paragraph is not a heading.
        """
        if paragraph.style is None:
            return None

        def _extract(s: str) -> int | None:
            s_lower = s.lower().strip()
            if s_lower == "title":
                return 1
            if "heading" in s_lower:
                m = re.search(r"\d+$", s.strip())
                return int(m.group()) if m else 2
            return None

        style = paragraph.style
        while style is not None:
            for val in (style.name or "", style.style_id or ""):
                level = _extract(val)
                if level is not None:
                    return max(1, min(level, 9))
            style = getattr(style, "base_style", None)
        return None

    def _is_caption(self, element: Any) -> bool:
        """Detect Word captions: they contain a SEQ field instruction."""
        for instr in element.findall(f".//{{{_W}}}instrText"):
            if instr.text and "SEQ" in instr.text:
                return True
        return False

    def _add_headers_footers(self) -> None:
        """Append header and footer text blocks for each document section."""
        doc = self._doc
        if doc is None:
            return

        odd_even = doc.settings.odd_and_even_pages_header_footer
        for section in doc.sections:
            hdrs = [section.header]
            if odd_even:
                hdrs.append(section.even_page_header)
            if section.different_first_page_header_footer:
                hdrs.append(section.first_page_header)

            seen_h: set[str] = set()
            for hdr in hdrs:
                text = " ".join(
                    p.text.strip() for p in hdr.paragraphs if p.text.strip()
                )
                if text and not text.isdigit() and text not in seen_h:
                    seen_h.add(text)
                    self._blocks.append({"type": "header", "content": text})

            ftrs = [section.footer]
            if odd_even:
                ftrs.append(section.even_page_footer)
            if section.different_first_page_header_footer:
                ftrs.append(section.first_page_footer)

            seen_f: set[str] = set()
            for ftr in ftrs:
                text = " ".join(
                    p.text.strip() for p in ftr.paragraphs if p.text.strip()
                )
                if text and not text.isdigit() and text not in seen_f:
                    seen_f.add(text)
                    self._blocks.append({"type": "footer", "content": text})
