from __future__ import annotations

import asyncio
import threading
import time
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from docx import Document

from comet_rag.engines.chunkers.base_chunker import RecursiveCharacterTextSplitter
from comet_rag.engines.chunkers.separators import SEPARATORS_MDX
from comet_rag.engines.chunkers.text_chunker import MdxChunker
from comet_rag.engines.cleaners.docx_cleaner import DocxCleaner
from comet_rag.engines.converters.archive_guard import (
    ArchiveLimits,
    ArchiveResourceLimitExceeded,
    validate_zip_archive,
)
from comet_rag.engines.converters.text_converter import DocxConverter
from comet_rag.engines.converters.types import DocxDocument
from comet_rag.engines.loaders.auto_loader import AutoLoader
from comet_rag.engines.loaders.base_loader import BaseLoader
from comet_rag.engines.loaders.types import LoaderContent, SourceContent, SourceType
from comet_rag.engines.loaders.url_loader import URLLoader
from comet_rag.engines.parsers.base_parser import BaseParser
from comet_rag.engines.parsers.docx_parser.docx_parser import DocxParser
from comet_rag.engines.parsers.types import DocxParsedContent


def test_heading_numbering_tracks_each_num_id_independently(monkeypatch) -> None:
    parser = DocxParser(heading_numbers=True)

    def formats(num_id: int) -> tuple[list[str], list[int]]:
        assert num_id in {1, 2}
        return ["%1"], [1]

    monkeypatch.setattr(parser, "_load_heading_formats", formats)

    assert parser._compute_heading_number(1, 1) == "1"
    assert parser._compute_heading_number(1, 1) == "2"
    assert parser._compute_heading_number(2, 1) == "1"
    assert parser._compute_heading_number(1, 1) == "3"


def test_empty_separators_are_rejected_at_construction() -> None:
    with pytest.raises(ValueError, match="separators 不能为空列表"):
        RecursiveCharacterTextSplitter(separators=[])


def test_directory_is_not_treated_as_local_file(tmp_path: Path) -> None:
    source = SourceContent(tmp_path)

    assert source.is_local is False
    assert source.pre_source_type is SourceType.UNKNOWN


class _RecordingLoader(BaseLoader):
    def load(self, source, *args, **kwargs):  # pragma: no cover
        raise NotImplementedError

    async def aload(self, source, *args, **kwargs):
        if isinstance(source, str):
            source = SourceContent(source)
        return LoaderContent(path=Path(source.source), source=source)

    def cleanup(self) -> None:
        pass


async def test_auto_loader_resolves_local_source_off_event_loop(
    tmp_path: Path, monkeypatch
) -> None:
    path = tmp_path / "doc.txt"
    path.write_text("content", encoding="utf-8")
    calls: list[object] = []

    async def recording_to_thread(func, *args):
        calls.append(func)
        return func(*args)

    monkeypatch.setattr(asyncio, "to_thread", recording_to_thread)
    loader = AutoLoader(loaders={SourceType.LOCAL: _RecordingLoader()})

    result = await loader.aload(SourceContent(path))

    assert result.path == path
    assert calls == [loader._resolve]


class _VisionModel:
    def __init__(self) -> None:
        self.mime_types: list[str] = []

    def describe(self, base64_data: str, media_type: str, **kwargs) -> str:
        self.mime_types.append(media_type)
        return "jpeg description"

    async def adescribe(self, base64_data: str, media_type: str, **kwargs) -> str:
        self.mime_types.append(media_type)
        return "jpeg description"


def _image_content(**overrides) -> DocxParsedContent:
    block = {
        "type": "image",
        "id": "rId1",
        "format": "jpg",
        "content": "aGVsbG8=",
        **overrides,
    }
    return DocxParsedContent(metadata={}, blocks=[block])


def test_docx_cleaner_uses_canonical_jpeg_mime_type() -> None:
    vision = _VisionModel()

    result = DocxCleaner(vision_model=vision).clean_to_markdown(_image_content())

    assert result == "jpeg description"
    assert vision.mime_types == ["image/jpeg"]


def test_docx_cleaner_degrades_unsupported_image_format_to_placeholder(
    tmp_path: Path,
) -> None:
    vision = _VisionModel()
    output_dir = tmp_path / "output"

    result = DocxCleaner(vision_model=vision).clean_to_markdown(
        _image_content(format="emf", alt_text="architecture"),
        output_dir=output_dir,
    )

    assert result == "[image: architecture]"
    assert vision.mime_types == []
    assert (output_dir / "result.md").read_text(encoding="utf-8") == result
    assert list((output_dir / "images").iterdir()) == []


@pytest.mark.parametrize(
    ("parse_content", "filename"),
    [
        (_image_content(), "../../outside"),
        (_image_content(id="../../outside"), "result"),
        (_image_content(format="../../png"), "result"),
    ],
)
def test_docx_cleaner_rejects_path_traversal(
    tmp_path: Path,
    parse_content: DocxParsedContent,
    filename: str,
) -> None:
    with pytest.raises(ValueError, match="非法|不支持"):
        DocxCleaner().clean_to_markdown(
            parse_content, output_dir=tmp_path / "output", filename=filename
        )
    assert not (tmp_path / "outside.md").exists()


async def test_docx_parser_implements_async_base_contract() -> None:
    document = Document()
    document.add_paragraph("hello")
    converted = DocxDocument(elements=document, metadata={"source": "unit"})
    parser = DocxParser()

    assert isinstance(parser, BaseParser)
    parsed = await parser.aparse(converted)
    assert parsed.text == "hello"


async def test_docx_parser_serializes_concurrent_aparse_calls(monkeypatch) -> None:
    first = Document()
    first.add_paragraph("first document")
    second = Document()
    second.add_paragraph("second document")
    parser = DocxParser()
    original_walk = parser._walk
    state_lock = threading.Lock()
    active = 0
    max_active = 0

    def monitored_walk(container) -> None:
        nonlocal active, max_active
        with state_lock:
            active += 1
            max_active = max(max_active, active)
        try:
            # 没有解析锁时，两次 to_thread 会在此窗口稳定重叠。
            time.sleep(0.05)
            original_walk(container)
        finally:
            with state_lock:
                active -= 1

    monkeypatch.setattr(parser, "_walk", monitored_walk)
    results = await asyncio.gather(
        parser.aparse(DocxDocument(elements=first, metadata={"id": "first"})),
        parser.aparse(DocxDocument(elements=second, metadata={"id": "second"})),
    )

    assert max_active == 1
    assert [result.metadata["id"] for result in results] == ["first", "second"]
    assert [result.text for result in results] == ["first document", "second document"]


@pytest.mark.parametrize("method", ["batch_load", "abatch_load"])
async def test_url_batch_rejects_non_positive_concurrency(
    tmp_path: Path, method: str
) -> None:
    loader = URLLoader(download_dir=tmp_path)
    try:
        with pytest.raises(ValueError, match="max_concurrency 必须大于 0"):
            result = getattr(loader, method)([], max_concurrency=0)
            if asyncio.iscoroutine(result):
                await result
    finally:
        await loader.aclose()


def test_mdx_fenced_code_body_is_not_split_at_internal_blank_line() -> None:
    code_body = "first_call()\n\nsecond_call()\n```"
    text = (
        "# Heading\n"
        + "intro " * 15
        + "\n```python\n"
        + code_body
        + "\n"
        + "outro " * 15
    )

    chunks = MdxChunker(chunk_size=100, chunk_overlap=0).chunk(text)

    assert SEPARATORS_MDX.index("\n```") < SEPARATORS_MDX.index("\n\n")
    assert any("first_call()\n\nsecond_call()" in chunk for chunk in chunks)


def test_docx_archive_is_rejected_before_python_docx_expands_it(
    tmp_path: Path,
) -> None:
    path = tmp_path / "bomb.docx"
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"x" * 100_000)

    content = LoaderContent(
        path=path,
        source=SourceContent(path),
        metadata={"file_type": "docx"},
    )
    converter = DocxConverter(
        content,
        archive_limits=ArchiveLimits(max_compression_ratio=5.0),
    )

    with pytest.raises(ArchiveResourceLimitExceeded, match="compression ratio"):
        converter.to_docx()


def test_docx_archive_bounds_xml_element_count(tmp_path: Path) -> None:
    path = tmp_path / "many-elements.docx"
    xml = b"<root>" + b"<cell/>" * 20 + b"</root>"
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", xml)

    content = LoaderContent(
        path=path,
        source=SourceContent(path),
        metadata={"file_type": "docx"},
    )
    converter = DocxConverter(
        content,
        archive_limits=ArchiveLimits(
            max_compression_ratio=1_000.0,
            max_xml_elements=5,
        ),
    )

    with pytest.raises(ArchiveResourceLimitExceeded, match="XML.*elements"):
        converter.to_docx()


def test_docx_archive_bounds_xml_tail_text(tmp_path: Path) -> None:
    path = tmp_path / "long-tail.docx"
    xml = b"<root><cell/>" + b"x" * 64 + b"</root>"
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", xml)

    with pytest.raises(ArchiveResourceLimitExceeded, match="text/tail"):
        validate_zip_archive(
            path,
            ArchiveLimits(
                max_compression_ratio=1_000.0,
                max_xml_text_chars=16,
            ),
        )
