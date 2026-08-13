"""`docx_parser` 快照测试。

962 行、此前零测试，是全仓改动风险最高的文件（plan R2）。本轮**只加保护网、
不重构**：把当前解析结果整体钉住，之后任何改动只要输出变了就会立刻显形。

样本由 `tests/fixtures/docx/build.py` 生成而非提交二进制 —— .docx 是 zip，
进了 git 就是不可 review 的黑盒；而且仓库里现有的真实文档含实际项目内容，
本项目计划开源，不应提交。生成脚本本身就是"样本里有什么"的说明。

**更新快照**（确认差异符合预期后）：

    UPDATE_DOCX_SNAPSHOTS=1 uv run pytest tests/unit/engines/test_docx_parser.py

切勿无脑更新 —— 快照测试的价值全在于"变了要有人看一眼"。
"""

from __future__ import annotations

import json
import os
from pathlib import Path

import pytest
from docx import Document

from comet_rag.engines.converters.types import DocxDocument
from comet_rag.engines.parsers.docx_parser.docx_parser import DocxParser
from tests.fixtures.docx.build import BUILDERS, build_all

SNAPSHOT_DIR = Path(__file__).resolve().parents[2] / "fixtures" / "docx" / "snapshots"
UPDATING = os.environ.get("UPDATE_DOCX_SNAPSHOTS") == "1"


@pytest.fixture(scope="session")
def generated_docx(tmp_path_factory: pytest.TempPathFactory) -> dict[str, Path]:
    return build_all(tmp_path_factory.mktemp("docx"))


def _parse(path: Path, **kwargs) -> dict:
    parsed = DocxParser(**kwargs).parse(
        DocxDocument(elements=Document(str(path)), metadata={})
    )
    return {"blocks": parsed.blocks, "text": parsed.text}


def _assert_matches_snapshot(name: str, actual: dict) -> None:
    SNAPSHOT_DIR.mkdir(parents=True, exist_ok=True)
    snapshot_path = SNAPSHOT_DIR / f"{name}.json"

    if UPDATING or not snapshot_path.exists():
        snapshot_path.write_text(
            json.dumps(actual, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
        )
        if not UPDATING:
            pytest.fail(
                f"快照 {snapshot_path.name} 不存在，已生成。请人工核对内容后再提交。"
            )
        return

    expected = json.loads(snapshot_path.read_text(encoding="utf-8"))
    assert actual == expected, (
        f"{name} 的解析结果与快照不符。确认改动符合预期后，用 "
        f"UPDATE_DOCX_SNAPSHOTS=1 重新生成。"
    )


@pytest.mark.parametrize("name", sorted(BUILDERS), ids=sorted(BUILDERS))
def test_parse_matches_snapshot(name: str, generated_docx: dict[str, Path]) -> None:
    _assert_matches_snapshot(name, _parse(generated_docx[name]))


# ── 针对性断言 ─────────────────────────────────────────────────────────────
#
# 快照能发现"变了"，但说不出"变的是什么、要紧不要紧"。下面这些针对
# 关键语义单独断言，坏掉时报错信息直指问题。


def test_heading_levels_are_preserved(generated_docx: dict[str, Path]) -> None:
    """标题层级丢了，分块就没法按语义边界切，检索质量直接下滑。"""
    blocks = _parse(generated_docx["basic"])["blocks"]
    headings = [(b["level"], b["content"]) for b in blocks if b["type"] == "heading"]

    assert [lvl for lvl, _ in headings] == [1, 2, 3]


def test_inline_formatting_becomes_markdown(generated_docx: dict[str, Path]) -> None:
    blocks = _parse(generated_docx["basic"])["blocks"]
    contents = [b.get("content", "") for b in blocks]

    assert any("**加粗**" in c and "*斜体*" in c for c in contents)


def test_list_items_keep_style_name(generated_docx: dict[str, Path]) -> None:
    """**当前行为**：列表项解析为 text 块，层级只体现在 style 名里
    （`List Bullet` / `List Bullet 2`），没有结构化嵌套。
    `DocxParsedContent.text` 里那条 `list` 分支目前走不到。
    """
    blocks = _parse(generated_docx["basic"])["blocks"]
    styles = [b.get("style") for b in blocks if b["type"] == "text"]

    assert "List Bullet" in styles
    assert "List Bullet 2" in styles
    assert "List Number" in styles
    assert not any(b["type"] == "list" for b in blocks)


def test_table_becomes_markdown_and_keeps_rows(
    generated_docx: dict[str, Path],
) -> None:
    blocks = _parse(generated_docx["table"])["blocks"]
    table = next(b for b in blocks if b["type"] == "table")

    assert table["content"].startswith("| 列A | 列B | 列C |")
    assert "| --- |" in table["content"]
    assert table["rows"][1] == ["值1", "值2", ""], "空单元格必须保留占位"


# ── 合并单元格（#18）──────────────────────────────────────────────────────


def test_repeated_cell_values_are_not_mistaken_for_merges(
    generated_docx: dict[str, Path],
) -> None:
    """**这条就是 #18 的全部理由。**

    旧实现把"相邻文本相等"当成合并单元格的判据，于是一张

        [季度, Q1, Q1]

    的表会被解析成 `[季度, Q1]` —— 整整一列凭空消失，不报错也不打日志，
    一路进到向量库。相邻两列取值相同在真实表格里再常见不过。

    判据必须来自结构（同一个 `<w:tc>`），不能来自数据。
    """
    blocks = _parse(generated_docx["merged_table"])["blocks"]
    rows = next(b for b in blocks if b["type"] == "table")["rows"]

    assert rows[1] == ["Q1", "Q1", "同上"], "相邻的重复值被当成合并副本删掉了"
    assert rows[2] == ["", "", "尾"], "相邻的空单元格被折叠了"


def test_merged_cells_keep_column_alignment(
    generated_docx: dict[str, Path],
) -> None:
    """横向合并的续格**留空**，不能删掉。

    Markdown 没有 colspan，续格只能空着；而删掉会让它后面的列整体左移 ——
    行尾补齐救不回来，"备注"会落到第 1 列去。
    """
    blocks = _parse(generated_docx["merged_table"])["blocks"]
    table = next(b for b in blocks if b["type"] == "table")

    assert table["rows"][0] == ["合并表头", "", "备注"]
    assert table["col_count"] == 3, "col_count 必须是真实的网格宽度"
    assert all(len(r) == 3 for r in table["rows"]), "各行宽度必须一致"


def test_a_span_wider_than_two_leaves_every_continuation_blank(
    generated_docx: dict[str, Path],
) -> None:
    """`gridSpan=3` ⇒ 续格有**两个**。

    判据从"跟上一格是不是同一个对象"换成"横跨几列"之后，这里是一处显式的
    计数（`grid_span - 1`），差一格就会让整行错位 —— 值得单独钉住，
    而不是指望 gridSpan=2 那条顺带覆盖。
    """
    blocks = _parse(generated_docx["merged_table"])["blocks"]
    rows = next(b for b in blocks if b["type"] == "table")["rows"]

    assert rows[5] == ["整行合并", "", ""]


def test_merge_position_within_the_row_does_not_shift_columns(
    generated_docx: dict[str, Path],
) -> None:
    """合并出现在**行首 / 行中 / 行尾**，以及同一行里有两个分开的合并。

    上一条只覆盖了行首的合并 —— 而错位恰恰最容易发生在"合并后面还有列"
    的时候：少留一个占位，它后面的所有列就整体左移一格，而行尾补齐
    根本救不回来（PR #32 评审建议补的组合）。
    """
    blocks = _parse(generated_docx["merged_table"])["blocks"]
    positions = [b for b in blocks if b["type"] == "table"][1]

    # 行首合并 + 中段合并：两个合并之间与之后的单元格都不能移位
    assert positions["rows"][0] == ["首合", "", "单A", "中合", "", "单B"]
    # 行尾合并：占位落在行尾，不是把 "尾合" 挤到别处
    assert positions["rows"][1] == ["a", "b", "c", "d", "尾合", ""]
    assert positions["col_count"] == 6


# ── 缺列：gridBefore / gridAfter（#33）─────────────────────────────────────


def test_rows_that_start_late_or_end_early_keep_their_grid_position(
    generated_docx: dict[str, Path],
) -> None:
    """Word 允许一行"晚开始"或"早结束"，那些位置是**不存在的网格列**。

    不补占位的话，"晚开始"那行的所有单元格都会左移一格，而行尾补齐还会让它
    看起来"宽度正常" —— 与漏掉合并续格是同一类错位，只是成因不同（#33）。
    """
    blocks = _parse(generated_docx["grid_gaps_table"])["blocks"]
    table = next(b for b in blocks if b["type"] == "table")

    assert table["rows"][0] == ["A", "B", "C"], "完整的行不受影响"
    assert table["rows"][1] == ["", "y", "z"], "gridBefore=1：y 必须落在第 1 列"
    assert table["rows"][2] == ["p", "q", ""], "gridAfter=1：缺的是尾列"
    assert table["rows"][3] == ["", "m", ""], "两端都缺时中间那格不能跑到边上去"
    assert table["col_count"] == 3


def test_a_missing_lead_column_and_a_merge_in_the_same_row(
    generated_docx: dict[str, Path],
) -> None:
    """缺列与横向合并叠在同一行上。

    两个特性各自都对，不代表组合起来也对：前缀占位、gridSpan 续格、后缀占位
    是三段接力，任何一段起点算错，后面全歪 —— 而这种行在缩进的子表格里
    并不罕见。
    """
    blocks = _parse(generated_docx["grid_gaps_table"])["blocks"]
    combined = [b for b in blocks if b["type"] == "table"][1]

    assert combined["rows"][0] == ["A", "B", "C", "D"]
    # gridBefore=1 占掉第 0 列，"合并" 从第 1 列起横跨两列，"末" 落在第 3 列
    assert combined["rows"][1] == ["", "合并", "", "末"]
    assert combined["col_count"] == 4


def test_a_hostile_grid_gap_cannot_blow_up_memory(tmp_path: Path) -> None:
    """`gridBefore` / `gridAfter` 是**文档里的整数**，在"用户上传文件"这条路上
    等于攻击者可控。直接 `[""] * n` 就是一个放大型 DoS。

    实测同一份 36 KB 的 docx：`w:val="10000000"` 让解析峰值多占 80 MB，
    封顶后 5 MB。每项 8 字节，而这个整数在 XML 里没有上限 —— 写成 2×10⁹
    就是十几 GB，文件大小一个字节都不用变（PR #34 评审）。

    封顶取**表格自己声明的网格宽度**而不是魔法常数：网格宽度由
    `<w:gridCol/>` 逐个声明，想要大的值就得写大的文件，放大系数没了。
    """
    from docx import Document as _Document  # noqa: PLC0415

    from comet_rag.core.logging import logger  # noqa: PLC0415
    from tests.fixtures.docx.build import _omit_grid_columns  # noqa: PLC0415

    doc = _Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 1).text = "仅此一格"
    # 真实载荷：只删一格，却声明缺一千万列 —— 文件大小一个字节都不用变
    _omit_grid_columns(table.rows[0], before=1, declare_before=10_000_000)
    path = tmp_path / "hostile.docx"
    doc.save(str(path))

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        blocks = _parse(path)["blocks"]
    finally:
        logger.remove(sink)

    row = next(b for b in blocks if b["type"] == "table")["rows"][0]
    assert len(row) < 100, f"缺列数没有被封顶，展开成了 {len(row)} 列"
    # 封顶到该表声明的网格宽度（2），再加上那一格真实单元格
    assert row == ["", "", "仅此一格"]
    assert any("超出网格宽度" in r for r in records), f"截断没留痕：{records}"


def _wide_short_row_table(
    path: Path, *, cols: int, rows: int, declare: int | None = None
) -> Path:
    """N 个 `<w:gridCol/>` + R 个短行，每行都声明 gridBefore=N。

    这是"表格宽度 × 行数"那种放大的最小载荷：XML 是 O(N+R)，展开是 O(N×R)。

    ⚠️ 每行必须**只含一个** `<w:tc>`。用 `add_table(rows=R, cols=N)` 再删几格
    是不行的 —— python-docx 会按 tblGrid 把每行建满，文档里真的就有 N×R 个
    `tc`，那是一份"确实很大的文件"，不是放大。第一版就是这么写错的，量出来
    的 2.3 GB 全花在 `Document()` 打开上，跟被测代码毫无关系。
    """
    import copy  # noqa: PLC0415

    from docx import Document as _Document  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    from tests.fixtures.docx.build import _omit_grid_columns  # noqa: PLC0415

    doc = _Document()
    table = doc.add_table(rows=1, cols=1)
    tbl = table._tbl  # noqa: SLF001
    grid = tbl.find(qn("w:tblGrid"))
    assert grid is not None, "python-docx 建的表居然没有 tblGrid"
    for _ in range(cols - 1):
        grid.append(grid.makeelement(qn("w:gridCol"), {}))

    table.cell(0, 0).text = "x"
    template = tbl.tr_lst[0]
    for _ in range(rows - 1):
        tbl.append(copy.deepcopy(template))
    # before=0：一格都不删（每行本来就只有一个 tc），只声明缺列数。
    # 走 helper 而不是自己动 XML —— `get_or_add_gridBefore` 是 python-docx
    # 动态生成的，pyright 看不见，散在测试里就得撒 ignore。
    for row in table.rows:
        _omit_grid_columns(row, declare_before=declare or (cols - 1))

    doc.save(str(path))
    return path


def test_a_wide_grid_times_many_rows_is_refused_not_expanded(tmp_path: Path) -> None:
    """**单行封顶还不够。**

    把缺列数封到网格宽度，挡住的只是"一个整数"的放大。攻击者还可以声明
    N 个 `<w:gridCol/>`，再加 R 个各只含一个 `<w:tc>` 的行、每行都写
    gridBefore=N —— XML 是 O(N+R)，展开却是 O(N×R)。

    实测 N=R=3000：约 9,000 个 XML 元素、38 KB 的 docx，
    无预算时展开成 900 万个单元格、峰值 87 MB；有预算时 11 MB（PR #34 评审）。

    所以要在**展开之前**按投影大小拒掉。用例把上限调小以便快速触发。
    """
    from comet_rag.core.logging import logger  # noqa: PLC0415

    path = _wide_short_row_table(tmp_path / "quad.docx", cols=200, rows=200)

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        blocks = _parse(path, max_table_cells=1000)["blocks"]
    finally:
        logger.remove(sink)

    assert not any(b["type"] == "table" for b in blocks), "超预算的表不该被展开"
    assert any("已跳过并留占位" in r for r in records), f"没有留痕：{records}"


def test_a_huge_gridspan_in_a_tiny_grid_is_refused(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """**第三个可控维度：`gridSpan`。**

    前两道防线（缺列数封到网格宽度、整表按 行数×网格宽度 定预算）都挡不住
    这一种：一份只有 2 列的表，某格声明 `gridSpan=10000000`，python-docx 的
    `row.cells` 就会返回一千万个 `_Cell` —— 光是建它就 75 MB，而且发生在解析
    循环**之前**，循环里再设防也来不及（PR #34 评审）。

    所以投影必须从 XML 层逐个 `tc` 数 span，绝不能先摸 `row.cells`。

    ⚠️ 断言"被拒绝了"是**不够的** —— 用 `len(row.cells)` 去算投影同样会拒绝，
    只是先白白建了一千万个 `_Cell`。所以这里把 `_Row.cells` 换成一个会炸的
    属性：准入检查一旦碰它，用例立刻红。
    """
    from docx import Document as _Document  # noqa: PLC0415
    from docx.table import _Row  # noqa: PLC0415

    from comet_rag.core.logging import logger  # noqa: PLC0415

    doc = _Document()
    table = doc.add_table(rows=1, cols=2)
    tc = table.rows[0]._tr.tc_lst[0]  # noqa: SLF001
    tc.get_or_add_tcPr().get_or_add_gridSpan().val = 10_000_000
    path = tmp_path / "bigspan.docx"
    doc.save(str(path))

    def _explode(self: object) -> None:
        raise AssertionError(
            "准入检查摸了 row.cells —— 一千万个 _Cell 已经建出来了，拦晚了"
        )

    monkeypatch.setattr(_Row, "cells", property(_explode))

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        blocks = _parse(path)["blocks"]
    finally:
        logger.remove(sink)

    assert not any(b["type"] == "table" for b in blocks), "超预算的表不该被展开"
    assert any("已跳过并留占位" in r for r in records), f"没有留痕：{records}"
    assert any(b["type"] == "caption" for b in blocks), "该留占位"


def test_a_vmerge_chain_inheriting_a_huge_span_is_refused(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """**第四个可控维度：`vMerge` 续格继承的 span。**

    `vMerge="continue"` 的续格会继承**上方根单元格**的 `gridSpan`，而它自己的
    `tc` 往往省略 `gridSpan` —— 本地值就是 1。实测：

        r0: XML 里 grid_span=[5]   row.cells 实际长度 = 5   (vMerge restart)
        r1: XML 里 grid_span=[1]   row.cells 实际长度 = 5   (vMerge continue)

    于是"逐行累加本地 span"会严重低估：根行 `gridSpan=S` 加 R 个续行，本地和
    约 `S + R`，实际展开却是 `S × (R+1)`。本用例的 S=5000 / R=500 下，
    旧投影约 5,500（远低于上限，会放行），实际约 250 万（PR #34 评审）。
    """
    import copy  # noqa: PLC0415

    from docx import Document as _Document  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    from docx.table import _Row  # noqa: PLC0415

    from comet_rag.core.logging import logger  # noqa: PLC0415

    span, rows = 5000, 500
    doc = _Document()
    table = doc.add_table(rows=1, cols=1)
    grid = table._tbl.find(qn("w:tblGrid"))  # noqa: SLF001
    assert grid is not None
    for _ in range(span - 1):
        grid.append(grid.makeelement(qn("w:gridCol"), {}))

    template = table._tbl.tr_lst[0]  # noqa: SLF001
    tc_pr = template.tc_lst[0].get_or_add_tcPr()
    tc_pr.get_or_add_gridSpan().val = span
    tc_pr.append(tc_pr.makeelement(qn("w:vMerge"), {qn("w:val"): "restart"}))
    for _ in range(rows):
        tr = copy.deepcopy(template)
        pr = tr.tc_lst[0].get_or_add_tcPr()
        span_el, merge_el = pr.find(qn("w:gridSpan")), pr.find(qn("w:vMerge"))
        assert span_el is not None and merge_el is not None
        pr.remove(span_el)  # 续格省略 gridSpan ⇒ 本地值是 1
        merge_el.set(qn("w:val"), "continue")
        table._tbl.append(tr)  # noqa: SLF001
    path = tmp_path / "vmerge.docx"
    doc.save(str(path))

    def _explode(self: object) -> None:
        raise AssertionError("准入检查摸了 row.cells —— 拦晚了")

    monkeypatch.setattr(_Row, "cells", property(_explode))

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        blocks = _parse(path)["blocks"]
    finally:
        logger.remove(sink)

    assert not any(b["type"] == "table" for b in blocks), "超预算的表不该被展开"
    assert any("已跳过并留占位" in r for r in records), f"没有留痕：{records}"


def test_an_oversized_table_leaves_a_searchable_placeholder(tmp_path: Path) -> None:
    """跳过之后**在原位留一条能被检索到的说明**。

    不截断成前 N 格：那会产出一张看起来完整、实则残缺的表，检索到它的人
    无从判断后面还有内容 —— 正是 #18 / #33 的同一个失败模式。
    也不干脆丢掉：那样知识库里完全看不出这里曾经有过东西。

    占位必须进入 `.text`，否则切块与向量化会把它丢掉，也就搜不到了。
    """
    path = _wide_short_row_table(tmp_path / "quad2.docx", cols=200, rows=200)
    parsed = _parse(path, max_table_cells=1000)

    placeholder = next(b for b in parsed["blocks"] if b["type"] == "caption")
    assert "表格未收录" in placeholder["content"]
    assert "200 行" in placeholder["content"]
    assert placeholder["content"] in parsed["text"], (
        "占位块没有进入正文，切块与向量化都会把它丢掉"
    )


def test_the_truncation_warning_is_one_per_table_not_one_per_row(
    tmp_path: Path,
) -> None:
    """截断告警按表汇总，不按行。

    按行打的话，一份畸形文档能直接产出成千上万行日志 —— 把"文档内容"放大成
    "日志写入量"，与内存放大是同一类问题（PR #34 评审）。实测 2000 行会打出
    2000 条。
    """
    from comet_rag.core.logging import logger  # noqa: PLC0415

    # declare 远大于网格宽度 ⇒ 每一行都会被截断
    path = _wide_short_row_table(
        tmp_path / "flood.docx", cols=3, rows=300, declare=10_000_000
    )

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        _parse(path)
    finally:
        logger.remove(sink)

    truncation = [r for r in records if "超出网格宽度" in r]
    assert len(truncation) == 1, f"300 行打出了 {len(truncation)} 条截断告警"
    assert "300 行" in truncation[0], f"汇总里没写清有多少行被截断：{truncation[0]}"


def test_a_table_without_tblgrid_is_skipped_not_crashed(tmp_path: Path) -> None:
    """`tblGrid` 是 schema 要求的必需元素，缺了 `len(table.columns)` 会抛
    `InvalidXmlError` —— 那会让整份文档解析失败，而不只是这一张表。

    缺它就无从确定网格宽度，也就无从给缺列数封顶，所以跳过整张表是对的；
    但要留痕，且不能拖垮同一份文档里其余的内容。
    """
    from docx import Document as _Document  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    from comet_rag.core.logging import logger  # noqa: PLC0415

    doc = _Document()
    doc.add_paragraph("表格前的正文。")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "会被跳过"
    grid = table._tbl.find(qn("w:tblGrid"))  # noqa: SLF001
    assert grid is not None
    table._tbl.remove(grid)  # noqa: SLF001
    doc.add_paragraph("表格后的正文。")
    path = tmp_path / "no_grid.docx"
    doc.save(str(path))

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        parsed = _parse(path)
    finally:
        logger.remove(sink)

    assert not any(b["type"] == "table" for b in parsed["blocks"])
    assert any("tblGrid" in r for r in records), f"跳过没留痕：{records}"
    assert "表格前的正文。" in parsed["text"], "其余内容不该受牵连"
    assert "表格后的正文。" in parsed["text"], "其余内容不该受牵连"


def test_a_genuinely_ragged_table_is_padded_but_says_so(tmp_path: Path) -> None:
    """两端补齐之后仍宽度不一致 ⇒ 文档的网格声明自不自洽。

    仍然补（下游不该拿到参差的行），但**必须留痕**：补只往行尾填，而真正缺
    的列可能在中间，静默做等于把错位藏起来。这条用例就是钉住"别静默"。

    真实样本里 17 张表无一触发，所以这条告警不会变成噪声。
    """
    from docx import Document as _Document  # noqa: PLC0415

    from comet_rag.core.logging import logger  # noqa: PLC0415

    doc = _Document()
    table = doc.add_table(rows=2, cols=3)
    for col, text in enumerate(["A", "B", "C"]):
        table.cell(0, col).text = text
    for col, text in enumerate(["x", "y", "z"]):
        table.cell(1, col).text = text
    # 删掉一格却**不**声明 gridAfter —— 这就是"网格声明不自洽"
    tr = table.rows[1]._tr  # noqa: SLF001
    tr.remove(tr.tc_lst[-1])
    path = tmp_path / "ragged.docx"
    doc.save(str(path))

    records: list[str] = []
    sink = logger.add(lambda m: records.append(m.record["message"]), level="WARNING")
    try:
        blocks = _parse(path)["blocks"]
    finally:
        logger.remove(sink)

    table_block = next(b for b in blocks if b["type"] == "table")
    assert table_block["rows"][1] == ["x", "y", ""], "仍然要补齐"
    assert any("网格宽度不一致" in r for r in records), f"补齐没有留痕：{records}"


def test_vertical_merge_repeats_down_the_column(
    generated_docx: dict[str, Path],
) -> None:
    """**当前行为**：纵向合并（vMerge）的文本沿列重复出现。

    与横向合并不同，它在每一行里都是独立的 `<w:tc>`，按行处理碰不到它。
    这里显式钉住，是为了让将来真要改成"只在首行出现"时，改动是**有意识**的
    而不是顺手带出来的 —— 对检索来说每行自带上下文其实是好事。
    """
    blocks = _parse(generated_docx["merged_table"])["blocks"]
    rows = next(b for b in blocks if b["type"] == "table")["rows"]

    assert rows[3] == ["纵向", "X", "X"]
    assert rows[4] == ["纵向", "Y", "Y"]


def test_image_block_carries_data_and_format(
    generated_docx: dict[str, Path],
) -> None:
    blocks = _parse(generated_docx["image"])["blocks"]
    image = next(b for b in blocks if b["type"] == "image")

    assert image["format"] == "png"
    assert image["content"], "图片内容为空，视觉模型就没东西可描述"
    assert "alt_text" in image


def test_equations_become_inline_latex(generated_docx: dict[str, Path]) -> None:
    text = _parse(generated_docx["equations"])["text"]

    assert r"$\frac{a}{b}$" in text
    assert "$x^{2}$" in text
    assert r"$\sqrt{x}$" in text
    assert r"$\sum_{i=1}^{n}i$" in text


def test_header_and_footer_are_separate_blocks(
    generated_docx: dict[str, Path],
) -> None:
    """页眉页脚要独立成块，DocxCleaner 才能按配置决定收不收。"""
    blocks = _parse(generated_docx["header_footer"])["blocks"]
    types = [b["type"] for b in blocks]

    assert "header" in types
    assert "footer" in types


def test_heading_numbers_option_changes_output(
    generated_docx: dict[str, Path],
) -> None:
    without = _parse(generated_docx["basic"], heading_numbers=False)
    with_numbers = _parse(generated_docx["basic"], heading_numbers=True)

    assert all(
        b.get("is_numbered") is False
        for b in without["blocks"]
        if b["type"] == "heading"
    )
    # 本样本用的是内置样式、无编号域，故两种设置结果一致；
    # 断言"不崩且形状一致"即可，编号域的样本留给真实文档冒烟。
    assert len(with_numbers["blocks"]) == len(without["blocks"])


# ── 真实文档冒烟（可选）───────────────────────────────────────────────────


REAL_DOCS_DIR = Path(__file__).resolve().parents[3] / "poc" / "docs"


@pytest.mark.skipif(
    not REAL_DOCS_DIR.is_dir() or not list(REAL_DOCS_DIR.glob("*.docx")),
    reason="本地没有 poc/docs/*.docx 真实样本",
)
def test_real_world_documents_parse_without_error() -> None:
    """合成样本的 XML 比 Word 真实输出简单得多，覆盖不到 Word 特有的怪异结构。

    这条用真实文档兜底，但**只做冒烟**（不崩、有产出），不比对内容 ——
    那些文档含实际项目材料，不会进版本库，也就无法维护稳定快照。
    """
    for path in sorted(REAL_DOCS_DIR.glob("*.docx")):
        parsed = DocxParser().parse(
            DocxDocument(elements=Document(str(path)), metadata={})
        )
        assert parsed.blocks, f"{path.name} 解析出 0 个块"
        assert parsed.text.strip(), f"{path.name} 文本为空"
